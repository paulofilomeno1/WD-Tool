import docx
import os


def createDocAll(data):
    cur_path = os.path.dirname(__file__)
    output_path = os.path.relpath("../output/services_and_products.docx", cur_path)
    print("Creating new Document...")
    my_doc = docx.Document()
    my_doc.add_heading("POOLS AND ACCESSORIES", 0)
    p = my_doc.add_paragraph("	The client by the name ")
    p.add_run(data["name"]).italic = True
    p.add_run(" required the service: ")
    p.add_run(data["service"]).bold = True
    p.add_run(" and " + str(data["quantity"]) + " " + data["product"] + " to be delivered.")
    for i in range(4):
        my_doc.add_paragraph(" ")
    my_doc.add_paragraph('Total Service Required:', style='Intense Quote')
    services = my_doc.add_paragraph("The service was required at: ")
    services.add_run(str(data["day"]) + "/" + str(data["month"]) + "/" + str(data["year"])).italic = True
    services.add_run(" by the client has a grant total of R$" + str(data["cost"]) + '.')
    for i in range(2):
        my_doc.add_paragraph(" ")
    my_doc.add_heading('The POOLS AND ACCESSORIES company appreciate the preference', level=1)
    my_doc.add_page_break()
    my_doc.save(output_path)
    print("Document created!")
